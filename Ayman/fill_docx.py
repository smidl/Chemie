"""Fill the SW-1 proposal template with the content of PROPOSAL.md.

Keeps the template's title page (its fields are filled in place), drops the pages of
"expected content" guidance, and writes the thirteen sections in their place using the
template's own styles so the result still looks like a faculty document.
"""
import re, copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "SW-1_Template Proposal_WS26+27 (1).docx"
OUT = "PROPOSAL_Ayman.docx"
MD  = "PROPOSAL.md"

doc = Document(SRC)
paras = doc.paragraphs

def settext(p, txt):
    """Replace a paragraph's text, keeping the first run's formatting."""
    if p.runs:
        p.runs[0].text = txt
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(txt)

# ---- 1. fill the title page in place
TITLE = ("The Effect of Noise-Model Assumptions on the Quality of Active Learning: "
         "A Sensitivity Study on Chemical Reaction Data")
for p in paras:
    t = p.text.strip()
    if t == "Master or Bachelor":
        settext(p, "Master")
    elif t == "Title of the Topic":
        settext(p, TITLE)
    elif t.startswith("in the winter / summer semester"):
        settext(p, "in the winter semester 2026/27")
    elif t.startswith("Prof. Dr. xxxx"):
        settext(p, "Prof. Dr. Václav Šmídl")
    elif t == "Submitted:":
        settext(p, "Submitted: August 2026")

# ---- 2. find where the guidance begins, and delete from there to the end of the body
start = None
for i, p in enumerate(paras):
    if p.text.strip().startswith("Expected content / structure"):
        start = i
        break
if start is None:
    for i, p in enumerate(paras):
        if p.text.strip().startswith("Hint:"):
            start = i
            break
anchor = paras[start]
for p in paras[start:]:
    p._element.getparent().remove(p._element)

# ---- 3. parse PROPOSAL.md into blocks
md = open(MD).read()
md = md.split("---\n", 1)[1] if "---\n" in md[:600] else md   # drop the front note

def clean(s):
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"\*(.+?)\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", s)
    return s.strip()

body = doc.element.body
def add(text, style=None, bold=False, size=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    if style:
        try: p.style = doc.styles[style]
        except KeyError: pass
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    if size: r.font.size = Pt(size)
    else: r.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

lines = md.split("\n")
i = 0
buf = []
def flush():
    global buf
    if buf:
        add(clean(" ".join(buf)))
        buf = []

in_table = False
# a markdown bullet may wrap over several lines; its continuations are indented, and must be
# joined back onto the bullet rather than emitted as separate paragraphs
pending_bullet = None

def flush_bullet():
    global pending_bullet
    if pending_bullet is None:
        return
    marker, txt = pending_bullet
    p = doc.add_paragraph()
    r = p.add_run(marker + clean(" ".join(txt)))
    r.font.name = "Arial"; r.font.size = Pt(11)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pending_bullet = None

while i < len(lines):
    ln = lines[i].rstrip()
    # continuation of a bullet: indented and not itself a new block
    if pending_bullet is not None and ln.startswith("  ") and ln.strip() \
       and not ln.lstrip().startswith(("- ", "|", "#", "*")) \
       and not re.match(r"^\s*\d+\. ", ln):
        pending_bullet[1].append(ln.strip()); i += 1; continue
    flush_bullet()
    if ln.startswith("## "):
        flush(); in_table = False
        add(clean(ln[3:]), bold=True, size=13, space_before=14, space_after=6)
    elif ln.startswith("*Alternatives") or ln.startswith("*Verified") or ln.startswith("*Follows"):
        flush()
        p = add(clean(ln)); p.runs[0].italic = True
    elif ln.startswith("|"):
        flush()
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if set("".join(cells)) <= set("-: "):
            i += 1; continue
        if not in_table:
            in_table = True
            add(clean(" · ".join(cells)), bold=True)
        else:
            add(clean("   ".join(cells)))
    elif ln.startswith("- ") or re.match(r"^\d+\. ", ln):
        flush(); in_table = False
        txt = clean(re.sub(r"^(- |\d+\. )", "", ln))
        # the template may not define list styles; fall back to a manual bullet/number
        marker = "\u2022  " if ln.startswith("- ") else re.match(r"^(\d+)\.", ln).group(1) + ".  "
        pending_bullet = (marker, [txt])
    elif not ln.strip():
        flush(); in_table = False
    else:
        if in_table: in_table = False
        buf.append(ln)
    i += 1
flush_bullet()
flush()

doc.save(OUT)
print("wrote", OUT)
